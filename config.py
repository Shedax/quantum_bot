import sqlite3
import docx
admin_chat_id=801093112
TOKEN=''

class DB:

    def __init__(self, database):
        self.connection = sqlite3.connect(database)
        self.cursor = self.connection.cursor()

    def select_all(self):
        """ Получаем все строки """
        with self.connection:
            return self.cursor.execute('SELECT * FROM lections').fetchall()

    def select_single(self, rownum):
        """ Получаем одну строку с номером rownum """
        with self.connection:
            return self.cursor.execute('SELECT * FROM words WHERE id = ?', (rownum,)).fetchone()

    def select_single_category(self, abbr):
        """ Получаем все строки с категорией abbr """
        with self.connection:
            return self.cursor.execute('SELECT * FROM words WHERE CATEGORY = ?', (abbr,)).fetchall()


    def select_lector_column(self, theme, type):
        """ Получаем все темы"""
        with self.connection:
            z = []
            for i in range(self.count_rows('users')):
                z += self.cursor.execute('SELECT lector FROM lections WHERE theme=? AND type=?',(theme, type)).fetchall()[i]
            return set(z)

    def count_rows(self, table):
        """ Считаем количество строк """
        with self.connection:
            result = self.cursor.execute('SELECT * FROM {}'.format(table)).fetchall()
            return len(result)
    def count_lector_rows(self, i):
        """ Считаем количество строк """
        with self.connection:
            return self.cursor.execute('SELECT * FROM lections WHERE lector=?', (i,)).fetchall()
    def get_vocabulary(self):
        """ Получаем все строки в словаре """
        with self.connection:
            result = self.cursor.execute('SELECT * FROM words WHERE VOC = 1').fetchall()
            return result

    def get_value(self, column, uid, current_table):
        """ Вставляем значения в БД """
        with self.connection:
            print(column)
            return self.cursor.execute('SELECT {0} FROM {1} WHERE user_id = {2}'.format(column, current_table, uid)).fetchone()

    def insert_value(self, column, value, uid, current_table):
        """ Вставляем значения в БД """
        with self.connection:
            sql = """
            UPDATE {0} 
            SET {1} = '%s'
            WHERE user_id = '%s'
            """.format(current_table, column) % (value, uid)
            self.cursor.execute(sql)

    def get_tests(self, current_field, num):
        """ Получаем все cлова из словаря"""
        with self.connection:
            tests=[]
            tests += self.cursor.execute('SELECT * FROM tests WHERE theme = ? AND num <= ?', (current_field, num)).fetchall()
            return tests

    def get_users(self, current_table):
        """ Получаем все cлова из словаря"""
        with self.connection:
            users=[]
            for i in range(self.count_rows(current_table)):
                users += self.cursor.execute('SELECT user_id FROM {}'.format(current_table)).fetchall()[i]
            return users

    def new_user(self, user_id, current_table):
        with self.connection:
            return self.cursor.execute('INSERT INTO {0}(user_id) VALUES ({1})'.format(current_table, user_id))

    def close(self):
        """ Закрываем текущее соединение с БД """
        self.connection.close()


class Doc:

    def __init__(self, document_name):
        self.document = docx.Document(document_name)

    def count_paragraphs(self):
        return len(self.document.paragraphs)

    def get_paragraph_text(self, i):
        return self.document.paragraphs[i].text
